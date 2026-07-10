"""Source loaders that normalise raw inputs into :class:`Document` objects.

Each loader fetches/reads a source (local path or URL), extracts plain text, and returns
one or more documents tagged with the appropriate :class:`SourceType`. Network access
uses ``httpx`` with ``tenacity`` retries; heavy parsing libraries are imported lazily.
"""

from __future__ import annotations

import os
from io import BytesIO
from urllib.parse import urlparse

from tenacity import retry, stop_after_attempt, wait_exponential

from app.schemas.rag import Document, SourceType

# Extension → SourceType mapping used for dispatch and inference.
_PROSE_EXTENSIONS: dict[str, SourceType] = {
    ".pdf": SourceType.PDF,
    ".docx": SourceType.DOCX,
    ".md": SourceType.MARKDOWN,
    ".markdown": SourceType.MARKDOWN,
    ".mdx": SourceType.MARKDOWN,
    ".html": SourceType.HTML,
    ".htm": SourceType.HTML,
    ".txt": SourceType.TXT,
    ".text": SourceType.TXT,
    ".rst": SourceType.TXT,
    ".log": SourceType.TXT,
}

_CODE_EXTENSIONS: set[str] = {
    ".py", ".js", ".jsx", ".mjs", ".cjs", ".ts", ".tsx", ".java", ".go", ".rs",
    ".rb", ".php", ".c", ".h", ".cpp", ".cc", ".cxx", ".hpp", ".cs", ".scala",
    ".swift", ".kt", ".kts", ".sol", ".lua", ".pl", ".hs", ".sh", ".bash",
    ".sql", ".r", ".m", ".dart", ".clj", ".ex", ".exs", ".vue", ".svelte",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".json", ".xml", ".gradle",
}

_HTTP_TIMEOUT = 30.0
_USER_AGENT = "ai-dev-assistant/1.0 (+ingestion)"


def _is_url(uri: str) -> bool:
    scheme = urlparse(uri).scheme
    return scheme in {"http", "https"}


def _extension(uri: str) -> str:
    path = urlparse(uri).path if _is_url(uri) else uri
    return os.path.splitext(path)[1].lower()


def infer_source_type(uri: str) -> SourceType:
    """Best-effort :class:`SourceType` from a URI's extension/scheme."""
    ext = _extension(uri)
    if ext in _PROSE_EXTENSIONS:
        return _PROSE_EXTENSIONS[ext]
    if ext in _CODE_EXTENSIONS:
        return SourceType.CODE
    if _is_url(uri):
        return SourceType.WEB
    return SourceType.TXT


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=0.5, min=0.5, max=8.0),
    reraise=True,
)
def _fetch_bytes(url: str) -> bytes:
    import httpx

    with httpx.Client(timeout=_HTTP_TIMEOUT, follow_redirects=True) as client:
        response = client.get(url, headers={"User-Agent": _USER_AGENT})
        response.raise_for_status()
        return response.content


def _read_bytes(uri: str) -> bytes:
    if _is_url(uri):
        return _fetch_bytes(uri)
    with open(uri, "rb") as handle:
        return handle.read()


def _read_text(uri: str) -> str:
    if _is_url(uri):
        return _fetch_bytes(uri).decode("utf-8", errors="replace")
    with open(uri, encoding="utf-8", errors="replace") as handle:
        return handle.read()


def _title_from_uri(uri: str) -> str:
    if _is_url(uri):
        parsed = urlparse(uri)
        name = os.path.basename(parsed.path) or parsed.netloc
        return name or uri
    return os.path.basename(uri) or uri


def load_pdf(uri: str, title: str | None = None) -> list[Document]:
    """Extract text from a PDF document (one document, page-delimited)."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(_read_bytes(uri)))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    content = "\n\n".join(pages)
    return [
        Document(
            content=content,
            source_type=SourceType.PDF,
            source_id=uri,
            title=title or _title_from_uri(uri),
            uri=uri,
            metadata={"pages": len(reader.pages)},
        )
    ]


def load_docx(uri: str, title: str | None = None) -> list[Document]:
    """Extract text from a DOCX document."""
    import docx

    document = docx.Document(BytesIO(_read_bytes(uri)))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    content = "\n".join(paragraphs)
    return [
        Document(
            content=content,
            source_type=SourceType.DOCX,
            source_id=uri,
            title=title or _title_from_uri(uri),
            uri=uri,
        )
    ]


def _html_to_text(raw: str) -> tuple[str, str | None]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    page_title = soup.title.string.strip() if soup.title and soup.title.string else None
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return cleaned, page_title


def load_html(uri: str, title: str | None = None) -> list[Document]:
    """Parse an HTML file/URL into plain text."""
    raw = _read_text(uri)
    content, page_title = _html_to_text(raw)
    return [
        Document(
            content=content,
            source_type=SourceType.HTML,
            source_id=uri,
            title=title or page_title or _title_from_uri(uri),
            uri=uri,
        )
    ]


def load_web(uri: str, title: str | None = None) -> list[Document]:
    """Fetch and extract readable text from a documentation web page."""
    raw = _read_text(uri)
    content, page_title = _html_to_text(raw)
    return [
        Document(
            content=content,
            source_type=SourceType.WEB,
            source_id=uri,
            title=title or page_title or _title_from_uri(uri),
            uri=uri,
        )
    ]


def load_markdown(uri: str, title: str | None = None) -> list[Document]:
    """Load a Markdown file as text."""
    return [
        Document(
            content=_read_text(uri),
            source_type=SourceType.MARKDOWN,
            source_id=uri,
            title=title or _title_from_uri(uri),
            uri=uri,
        )
    ]


def load_text(uri: str, title: str | None = None) -> list[Document]:
    """Load a plain-text file."""
    return [
        Document(
            content=_read_text(uri),
            source_type=SourceType.TXT,
            source_id=uri,
            title=title or _title_from_uri(uri),
            uri=uri,
        )
    ]


def load_code(uri: str, title: str | None = None) -> list[Document]:
    """Load a source-code file, tagging the inferred language."""
    ext = _extension(uri)
    return [
        Document(
            content=_read_text(uri),
            source_type=SourceType.CODE,
            source_id=uri,
            title=title or _title_from_uri(uri),
            uri=uri,
            metadata={"language": ext.lstrip(".")} if ext else {},
        )
    ]


def load_by_uri(
    uri: str, source_type: SourceType | None = None, title: str | None = None
) -> list[Document]:
    """Dispatch ``uri`` to the appropriate loader based on type/extension/scheme."""
    resolved = source_type or infer_source_type(uri)
    if resolved == SourceType.PDF:
        return load_pdf(uri, title)
    if resolved == SourceType.DOCX:
        return load_docx(uri, title)
    if resolved == SourceType.HTML:
        return load_html(uri, title)
    if resolved == SourceType.WEB:
        return load_web(uri, title)
    if resolved == SourceType.MARKDOWN:
        return load_markdown(uri, title)
    if resolved == SourceType.CODE:
        return load_code(uri, title)
    return load_text(uri, title)
